import io
import importlib
from docx import Document
from contextlib import redirect_stdout
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class HelpDocumentGenerator:
    """
    A class to generate a Word document from a list of Python objects such as modules, classes, and functions.
    """

    def __init__(self, output_file="reference_manual.docx"):
        """
        Initialize the generator with an output file for the Word document.
        
        :param output_file: Optional output filename for the Word document.
        """
        self.output_file = output_file
        self.doc = Document()  # Initialize the Word document
        self.toc_entries = []   # Store TOC entries

    def add_help_section(self, title, content):
        """
        Add a section to the Word document with a title and content.

        :param title: The title of the section (e.g., module or class name).
        :param content: The help text or other information to add to the document.
        """
        # Add the TOC entry
        self.toc_entries.append(title)
        self.doc.add_page_break()  # Start new section on a new page

        # Add heading and content
        self.doc.add_heading(title, level=1)
        self.doc.add_paragraph(content)

    def generate_document(self, package_data):
        """
        Generate the Word document by iterating through the package data and adding help content.

        :param package_data: Dictionary containing the introspected package data (modules, classes, functions, subpackages).
        """
        self._create_toc()

        # Add help content
        for modname in package_data['modules']:
            help_text = self._get_help_text(modname)
            self.add_help_section(f"Module: {modname}", help_text)

        for cls in package_data['classes']:
            help_text = self._get_help_text(cls)
            self.add_help_section(f"Class: {cls}", help_text)

        for func in package_data['functions']:
            help_text = self._get_help_text(func)
            self.add_help_section(f"Function: {func}", help_text)

        for subpkg in package_data['subpackages']:
            help_text = self._get_help_text(subpkg)
            self.add_help_section(f"Subpackage: {subpkg}", help_text)

        # Save the final document
        self.doc.save(self.output_file)
        print(f"Word document generated: {self.output_file}")

    def _create_toc(self):
        """
        Manually create a table of contents using the stored TOC entries.
        """
        self.doc.add_heading('Table of Contents', level=1)

        # For each TOC entry, add it as a hyperlink (not functional in docx, but we'll add a TOC field)
        for toc_entry in self.toc_entries:
            self.doc.add_paragraph(toc_entry)

        # Add a placeholder for an auto-generated TOC by Word
        toc_placeholder = self.doc.add_paragraph()
        run = toc_placeholder.add_run()
        fldChar = OxmlElement('w:fldChar')  # Create a field character element
        fldChar.set(qn('w:fldCharType'), 'begin')  # Specify it's the beginning of the field
        run._r.append(fldChar)

        instrText = OxmlElement('w:instrText')  # Create the instruction text element
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # TOC field instructions
        run._r.append(instrText)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')  # End of the field
        run._r.append(fldChar)

    @staticmethod
    def _get_help_text(obj_name):
        """
        Dynamically retrieve the full help text for an object by its name using the help() function.

        :param obj_name: The fully qualified name of the object.
        :return: The captured help text for the object.
        """
        try:
            obj = importlib.import_module(obj_name)
        except ImportError:
            return f"Could not import {obj_name}"

        # Capture the output of help() using StringIO
        help_text_io = io.StringIO()
        with redirect_stdout(help_text_io):
            help(obj)
        help_text = help_text_io.getvalue()

        return help_text
