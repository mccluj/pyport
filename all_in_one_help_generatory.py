import pkgutil
import importlib
import inspect
import io
from contextlib import redirect_stdout
from docx import Document

class PackageInspector:
    """
    A class to inspect a Python package and list its modules, classes, functions, and subpackages.
    """

    def __init__(self, package):
        """
        Initialize the inspector with a package and create containers for storing results.
        
        :param package: The root package to inspect (e.g., `my_package`).
        """
        self.package = package
        self.data = {
            'modules': [],
            'classes': [],
            'functions': [],
            'subpackages': []
        }

    def inspect(self):
        """
        Main method to inspect the package and gather its modules, classes, functions, and subpackages.
        """
        for importer, modname, ispkg in pkgutil.walk_packages(self.package.__path__, self.package.__name__ + "."):
            self._process_module(modname, ispkg)

        return self.data

    def _process_module(self, modname, ispkg):
        """
        Process each module or subpackage found in the package.
        
        :param modname: Name of the module or subpackage.
        :param ispkg: Boolean indicating if it's a package (True) or a module (False).
        """
        try:
            # Dynamically import the module or package
            module = importlib.import_module(modname)

            if ispkg:
                self.data['subpackages'].append(modname)
            else:
                self.data['modules'].append(modname)

            self._extract_members(module, modname)

        except ImportError as e:
            print(f"Error importing {modname}: {e}")

    def _extract_members(self, module, modname):
        """
        Extract classes and functions from the module and store them.
        
        :param module: The module object to inspect.
        :param modname: The module name for filtering its members.
        """
        for name, obj in inspect.getmembers(module):
            if inspect.isclass(obj) and obj.__module__ == modname:
                self.data['classes'].append(f"{modname}.{name}")
            elif inspect.isfunction(obj) and obj.__module__ == modname:
                self.data['functions'].append(f"{modname}.{name}")


class HelpDocumentGenerator:
    """
    A class to generate a Word document from a list of Python objects such as modules, classes, and functions.
    """

    def __init__(self, output_file="reference_manual.docx"):
        """
        Initialize the generator with an output file for the Word document.
        
        :param output_file: Optional output filename for the Word document.
        """
        self.output_file = output_file
        self.doc = Document()  # Initialize the Word document

    def add_help_section(self, title, content):
        """
        Add a section to the Word document with a title and content.

        :param title: The title of the section (e.g., module or class name).
        :param content: The help text or other information to add to the document.
        """
        self.doc.add_page_break()  # Start new section on a new page
        self.doc.add_heading(title, level=1)
        self.doc.add_paragraph(content)

    def generate_document(self, package_data):
        """
        Generate the Word document by iterating through the package data and adding help content.

        :param package_data: Dictionary containing the introspected package data (modules, classes, functions, subpackages).
        """
        for modname in package_data['modules']:
            help_text = self._get_help_text(modname)
            self.add_help_section(f"Module: {modname}", help_text)

        for cls in package_data['classes']:
            help_text = self._get_help_text(cls)
            self.add_help_section(f"Class: {cls}", help_text)

        for func in package_data['functions']:
            help_text = self._get_help_text(func)
            self.add_help_section(f"Function: {func}", help_text)

        for subpkg in package_data['subpackages']:
            help_text = self._get_help_text(subpkg)
            self.add_help_section(f"Subpackage: {subpkg}", help_text)

        self.doc.save(self.output_file)
        print(f"Word document generated: {self.output_file}")

    @staticmethod
    def _get_help_text(obj_name):
        """
        Dynamically retrieve the full help text for an object by its name using the help() function.

        :param obj_name: The fully qualified name of the object.
        :return: The captured help text for the object.
        """
        try:
            obj = importlib.import_module(obj_name)
        except ImportError:
            return f"Could not import {obj_name}"

        # Capture the output of help() using StringIO
        help_text_io = io.StringIO()
        with redirect_stdout(help_text_io):
            help(obj)
        help_text = help_text_io.getvalue()

        return help_text


# Main program to combine the package inspection and Word document generation
import my_package  # Replace with your package name

def main():
    # Step 1: Inspect the package to gather data
    inspector = PackageInspector(my_package)
    package_data = inspector.inspect()

    # Step 2: Generate the Word document using the gathered package data
