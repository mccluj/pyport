import io
from docx import Document
from contextlib import redirect_stdout

class HelpDocumentGenerator:
    """
    A class to generate a Word document with help() output for Python objects
    such as classes, modules, or functions.
    """

    # Constants for document formatting
    TITLE_LEVEL = 1  # Title heading level for each help section
    OUTPUT_FILE = "help_reference_manual.docx"  # Default output filename

    def __init__(self, objects, output_file=None):
        """
        Initialize the generator with a list of objects and optional output file.

        :param objects: List of objects (e.g., classes or modules) to document.
        :param output_file: Optional output filename for the Word document.
        """
        self.objects = objects
        self.output_file = output_file or self.OUTPUT_FILE
        self.doc = Document()  # Initialize the Word document

    @staticmethod
    def capture_help_text(obj):
        """
        Capture the help() output for a given object as a string.

        :param obj: The Python object to get help on.
        :return: The captured help text as a string.
        """
        help_text_io = io.StringIO()
        with redirect_stdout(help_text_io):
            help(obj)
        return help_text_io.getvalue()

    def add_help_section(self, obj):
        """
        Add a help section to the Word document for the given object.

        :param obj: The Python object to add help information for.
        """
        title = f"Help for {obj.__name__}"
        help_text = self.capture_help_text(obj)

        self.doc.add_page_break()  # Ensure new section starts on a new page
        self.doc.add_heading(title, level=self.TITLE_LEVEL)
        self.doc.add_paragraph(help_text)

    def generate_document(self):
        """
        Generate the Word document by iterating through the objects and
        saving the output to a file.
        """
        for obj in self.objects:
            self.add_help_section(obj)

        self.save_document()

    def save_document(self):
        """
        Save the Word document to the specified output file.
        """
        self.doc.save(self.output_file)
        print(f"Word document generated: {self.output_file}")


# Example usage with classes
class AssetClass1:
    """
    AssetClass1 manages some asset operations.

    Example:
    >>> ac = AssetClass1()
    >>> ac.method1(10)
    """

    def method1(self, x):
        """
        Multiplies input by 2.

        Example:
        >>> AssetClass1().method1(5)
        10
        """
        return x * 2


class AssetClass2:
    """
    AssetClass2 deals with another type of asset.

    Example:
    >>> ac = AssetClass2()
    >>> ac.method2(20)
    """

    def method2(self, y):
        """
        Adds 20 to the input.

        Example:
        >>> AssetClass2().method2(10)
        30
        """
        return y + 20


# List of asset classes or other Python objects to document
asset_classes = [AssetClass1, AssetClass2]

# Generate Word document with help documentation for the given classes
generator = HelpDocumentGenerator(asset_classes)
generator.generate_document()
